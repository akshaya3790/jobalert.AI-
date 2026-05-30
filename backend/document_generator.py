import os
import uuid
import asyncio
from jinja2 import Template
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HTML_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        body { font-family: "Helvetica Neue", Helvetica, Arial, sans-serif; font-size: 11pt; line-height: 1.4; color: #1f2937; margin: 0; padding: 40px; }
        h1 { font-size: 20pt; text-align: center; margin: 0 0 10px 0; color: #111827; }
        .contact { text-align: center; font-size: 10pt; color: #4b5563; margin-bottom: 20px; }
        h2 { font-size: 13pt; text-transform: uppercase; border-bottom: 1px solid #111827; margin-top: 20px; padding-bottom: 4px; color: #111827; }
        .summary { margin-bottom: 20px; }
        .job { margin-bottom: 16px; }
        .job-header { display: flex; justify-content: space-between; font-weight: bold; margin-bottom: 4px; }
        .job-meta { display: flex; justify-content: space-between; font-style: italic; color: #4b5563; margin-bottom: 6px; }
        ul { margin: 0; padding-left: 20px; }
        li { margin-bottom: 4px; text-align: justify; }
        .skills { margin-bottom: 20px; }
    </style>
</head>
<body>
    <h1>{{ resume.name | default('First Last') }}</h1>
    <div class="contact">
        {{ resume.email | default('email@example.com') }} | {{ resume.phone | default('') }} | {{ resume.location | default('') }}
    </div>

    {% if resume.summary %}
    <h2>Professional Summary</h2>
    <div class="summary">
        {{ resume.summary }}
    </div>
    {% endif %}

    <h2>Experience</h2>
    {% for job in resume.experience %}
    <div class="job">
        <div class="job-header">
            <span>{{ job.title }} | {{ job.company }}</span>
            <span>{{ job.start_date }} - {{ job.end_date | default('Present') }}</span>
        </div>
        <ul>
            {% for point in job.bullets %}
            <li>{{ point }}</li>
            {% endfor %}
        </ul>
    </div>
    {% endfor %}

    <h2>Education</h2>
    {% for edu in resume.education %}
    <div class="job">
        <div class="job-header">
            <span>{{ edu.degree }} - {{ edu.institution }}</span>
            <span>{{ edu.graduation_year }}</span>
        </div>
    </div>
    {% endfor %}

    <h2>Skills</h2>
    <div class="skills">
        {% if resume.skills %}
            {{ resume.skills | join(', ') }}
        {% endif %}
    </div>
</body>
</html>
"""

async def generate_pdf(resume_data: dict, output_path: str):
    """Generates an ATS-friendly PDF using Playwright."""
    template = Template(HTML_TEMPLATE)
    html_content = template.render(resume=resume_data)
    
    # Needs async playwright
    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=True)
        page = await browser.new_page()
        await page.set_content(html_content, wait_until="networkidle")
        await page.pdf(
            path=output_path, 
            format="Letter", 
            margin={"top": "0.5in", "bottom": "0.5in", "left": "0.5in", "right": "0.5in"},
            print_background=True
        )
        await browser.close()
        
    return output_path

def generate_docx(resume_data: dict, output_path: str):
    """Generates an ATS-friendly DOCX natively using python-docx."""
    document = Document()
    
    # Configure default style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Name Header
    name_para = document.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(resume_data.get('name', 'First Last'))
    name_run.bold = True
    name_run.font.size = Pt(20)

    # Contact Info
    contact_para = document.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_details = [
        resume_data.get('email', ''),
        resume_data.get('phone', ''),
        resume_data.get('location', '')
    ]
    contact_para.add_run(" | ".join([c for c in contact_details if c]))

    # Summary
    if resume_data.get('summary'):
        doc_add_heading(document, 'Professional Summary')
        document.add_paragraph(resume_data.get('summary'))

    # Experience
    if resume_data.get('experience'):
        doc_add_heading(document, 'Experience')
        for job in resume_data.get('experience', []):
            job_para = document.add_paragraph()
            title = job_para.add_run(f"{job.get('title', '')} | {job.get('company', '')}")
            title.bold = True
            
            date = job_para.add_run(f"    ({job.get('start_date', '')} - {job.get('end_date', 'Present')})")
            date.italic = True
            
            for bullet in job.get('bullets', []):
                document.add_paragraph(bullet, style='List Bullet')

    # Education
    if resume_data.get('education'):
        doc_add_heading(document, 'Education')
        for edu in resume_data.get('education', []):
            edu_para = document.add_paragraph()
            edu_run = edu_para.add_run(f"{edu.get('degree', '')} - {edu.get('institution', '')}")
            edu_run.bold = True
            document.add_paragraph(f"Graduation: {edu.get('graduation_year', '')}")

    # Skills
    if resume_data.get('skills'):
        doc_add_heading(document, 'Skills')
        skills = resume_data.get('skills', [])
        if isinstance(skills, list):
            skills = ", ".join(skills)
        document.add_paragraph(skills)

    document.save(output_path)
    return output_path

def doc_add_heading(document, text):
    heading = document.add_heading(level=2)
    run = heading.add_run(text.upper())
    run.font.size = Pt(13)
    run.bold = True

COVER_LETTER_HTML_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        body { font-family: "Helvetica Neue", Helvetica, Arial, sans-serif; font-size: 11pt; line-height: 1.6; color: #1f2937; margin: 0; padding: 40px; }
        .body-text { white-space: pre-wrap; margin-top: 20px; }
    </style>
</head>
<body>
    <div class="body-text">{{ letter_text }}</div>
</body>
</html>
"""

async def generate_cover_letter_pdf(letter_text: str, output_path: str):
    """Generates a PDF for a cover letter using Playwright."""
    template = Template(COVER_LETTER_HTML_TEMPLATE)
    html_content = template.render(letter_text=letter_text)
    
    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=True)
        page = await browser.new_page()
        await page.set_content(html_content, wait_until="networkidle")
        await page.pdf(
            path=output_path, 
            format="Letter", 
            margin={"top": "1in", "bottom": "1in", "left": "1in", "right": "1in"},
            print_background=True
        )
        await browser.close()
        
    return output_path

def generate_cover_letter_docx(letter_text: str, output_path: str):
    """Generates a DOCX for a cover letter using python-docx."""
    document = Document()
    
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    for paragraph in letter_text.split('\n'):
        if paragraph.strip():
            document.add_paragraph(paragraph.strip())
        else:
            document.add_paragraph() # preserve line breaks

    document.save(output_path)
    return output_path
